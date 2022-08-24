from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
import time
import requests
from io import StringIO
import io
from PIL import Image
import docx
from docx.shared import Inches
import os
from webdriver_manager.chrome import ChromeDriverManager
from os.path import dirname, join


# service = Service(r"C:\Users\Stefan\chromedriver\chromedriver")
# driver = webdriver.Chrome(service=service)
#
#
driver = webdriver.Chrome(ChromeDriverManager().install())
#
#
#
time.sleep(2)
driver.get("https://www.houseplans.net/")
driver.maximize_window()
time.sleep(2)
# name_of_the_house = input('What is the number of the house you are looking for?: ')
# 7922-00123

#


""" ABOUT ELEMENT """
def about():
    about_plan = driver.find_elements(By.CLASS_NAME, 'left.heading')[2]
    text_about = driver.find_element(By.XPATH, '//*[@id="wrapper"]/div[4]/div/div/div[5]/div/div[3]/p')
    doc.add_heading(about_plan.text)
    doc.add_paragraph(text_about.text)





""" FEATURES AND DETAILS ELEMENTS"""
def main_program():
    doc.add_heading(name_of_the_house)

    list_of_all_details = []
    features_element = driver.find_element(By.XPATH, '//*[@id="wrapper"]/div[4]/div/div/div[3]/div/div[4]/aside[1]/div/h2/small')
    features = features_element.text
    plan_details_elements = driver.find_elements(By.CLASS_NAME, 'plan_details')
    for i in plan_details_elements:
        details = 'DETAILS'
        if i.text[0:4] in 'TOTAL':
            list_of_all_details.append(details)
        list_of_all_details.append(i.text)

    list_of_all_details_without_colon = [o.replace(':\n', ': ') for o in list_of_all_details]
    list_of_all_details_without_space_sign = [z.replace('\n', ', ') for z in list_of_all_details_without_colon]

    doc.add_heading(features)
    for list in list_of_all_details_without_space_sign:
        if list[0:6] in 'DETAILS':
            doc.add_heading(list)
        else:
            doc.add_paragraph(list)

    including = driver.find_element(By.CLASS_NAME, 'faq-block')
    doc.add_paragraph(including.text)



""" MAIN PROGRAM """

scratch = True
while scratch == True:
    doc = docx.Document()
    name_of_the_house = input('What is the number of the house you are looking for?: ')
    number_of_house = driver.find_element(By.ID, 'find-by-plan-input')
    number_of_house.send_keys(name_of_the_house)
    search_button = driver.find_element(By.CLASS_NAME, 'fa.fa-search')
    search_button.click()
    time.sleep(5)
    main_program()
    about()

    """ PICTURE ELEMENTS """
    pictures_list = []
    pictures = driver.find_elements(By.CLASS_NAME, 'gallery-item')
    for picture in pictures:
        picture_link = picture.get_attribute('href')
        pictures_list.append(picture_link)

    """ if there are duplicates it will be deleted """
    pictures_simplified = list(set(pictures_list))

    for i in pictures_simplified:
        index = pictures_simplified.index(i)
        image_content = requests.get(i).content
        image_file = io.BytesIO(image_content)
        image = Image.open(image_file)
        image.save(f'image{index}.jpg')
        doc.add_picture(f'image{index}.jpg', width=Inches(5), height=Inches(4))


    """ SAVING WORD DOCUMENT """
    doc.save(f'{name_of_the_house}.docx')
    question = input('Do you want to scratch again new house number? (yes/no)')
    if question == 'yes':
        scratch = True
    else:
        scratch = False


""" DELETE PICTURES FROM FOLDER """

# """1. If we want to have project directory with RELATIVE path, not hard typed root (absolute path) """
# project_root = os.path.dirname(os.path.realpath(__file__))
#
# """ This is no need for this situation """
# output_path = join(project_root, 'dist')
#
#
#
# """ this is hard typed root (ABSOLUTE path) """
# # path = "C:\\Users\\Stefan\\Desktop\\house_plans"
#
#
# """2. using listdir() method to list the files of the folder """
# test = os.listdir(output_path)
#
# for images in test:
#     if images.endswith(".jpeg") or images.endswith('.jpg') or images.endswith('.png'):
#         os.remove(os.path.join(output_path, images))

